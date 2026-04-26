"""Тесты этапа 9А.2.5.

Покрывают:
  A. Reoptimize UI — короткий toast + компактная модалка с карточками
     (вместо «простыни» в toast'е). Авто-обновление UI после применения
     или отката (через location.reload + sessionStorage для scroll
     position).
  B. Word KP — нормализация полей страницы, NBSP в разрядных
     разделителях чисел, центрирование внешней таблицы.
"""

from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

import docx as _docx
import pytest


# =====================================================================
# A. Reoptimize UI — статический разбор project.js и шаблона
# =====================================================================
#
# Полноценно прокликать модалку без headless-браузера в pytest сложно,
# поэтому проверяем JS-источник: что он содержит нужные конструкции
# (короткий toast, модалку с card, текст «Состав и цены не изменились»,
# вызов location.reload после отката).

@pytest.fixture(scope="session")
def project_js() -> str:
    return Path("static/js/project.js").read_text(encoding="utf-8")


def test_reoptimize_response_has_minimal_toast(project_js):
    """Tост после reoptimize — короткая строка вида «Пересобрано N
    конфигурац…», без полного diff'а компонентов внутри toast'а.

    Проверяем отсутствие старой большой сборки `lines.push(deltaSummaryHtml)`
    внутри toast и наличие новой короткой формулировки.
    """
    # Новая формулировка короткого toast'а.
    assert "'Пересобрано '" in project_js
    # Toast больше НЕ собирает многоблочный HTML с deltaSummaryHtml
    # внутри — теперь deltaSummaryHtml/changeListHtml идут только в
    # модалку. Проверим, что в коде нет старой конструкции
    # «lines.push(deltaSummaryHtml(...))».
    assert "lines.push(" not in project_js, (
        "Старая модель «огромного toast'а с lines.push» должна быть "
        "удалена в пользу модалки."
    )


def test_reoptimize_modal_renders_when_changes(project_js):
    """JS создаёт модалку с class='modal-container' (общий класс
    дизайн-системы) и заголовком «Результат пересборки»."""
    assert "modal-container" in project_js
    assert "Результат пересборки" in project_js
    # Модалка строится в DOM динамически через ensureReoptimizeModal.
    assert "reoptimize-modal" in project_js
    # Есть кнопка «Применить».
    assert "Применить" in project_js


def test_reoptimize_no_changes_shows_card(project_js):
    """Карточка модалки для status='no_changes' содержит текст
    «Состав и цены не изменились»."""
    assert "Состав и цены не изменились" in project_js


def test_rollback_reloads_variants_block(project_js):
    """После применения / отката модалка перезагружает страницу
    (AJAX-фрагмента для variants нет, поэтому reload + сохранение
    scroll position через sessionStorage)."""
    assert "location.reload()" in project_js
    # sessionStorage используется для восстановления scroll-позиции.
    assert "sessionStorage" in project_js
    assert "SCROLL_STORAGE_KEY" in project_js
    # Toast про отмену пересборки перенесён в deferred-механизм
    # (показывается на следующей загрузке страницы).
    assert "Пересборка отменена" in project_js


# =====================================================================
# B. Word KP — нормализация полей, NBSP, центрирование таблицы
# =====================================================================

from app.services.export import kp_builder

_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _fake_items(items_spec):
    out = []
    for i, (uu, qty, name) in enumerate(items_spec, start=1):
        out.append({
            "id": i, "query_id": 100 + i, "variant_manufacturer": "Intel",
            "quantity": qty, "position": i,
            "auto_name": name, "custom_name": None, "display_name": name,
            "unit_usd": uu, "unit_rub": 0.0,
            "total_usd": round(uu * qty, 2), "total_rub": 0.0,
            "created_at": None, "updated_at": None,
        })
    return out


def _mock_rate(value: str = "90"):
    return patch(
        "app.services.export.kp_builder.exchange_rate.get_usd_rate",
        return_value=(Decimal(value), date(2026, 4, 25), "cache"),
    )


def _build_kp(items_spec, markup=15):
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=_fake_items(items_spec),
    ), _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=markup, db=None,
        )
    return _docx.Document(BytesIO(data))


def test_kp_page_margins_normalized():
    """После генерации в каждой секции pgMar = top=2cm, bottom=2cm,
    left=2cm, right=1.5cm. Без этого исходный шаблон даёт ассиметричные
    поля (top=284, right=140, bottom=0, left=1701 twips), из-за чего
    таблица сдвинута влево, а правое поле почти отсутствует.
    """
    from docx.shared import Cm
    doc = _build_kp([(100.0, 1, "Поля")])
    assert doc.sections, "В документе нет ни одной секции"
    for section in doc.sections:
        # python-docx возвращает Length (EMU). Переводим в cm с точностью
        # до сотых, потому что округления Cm() могут «гулять» на 1 EMU.
        top = round(float(section.top_margin.cm), 2)
        bottom = round(float(section.bottom_margin.cm), 2)
        left = round(float(section.left_margin.cm), 2)
        right = round(float(section.right_margin.cm), 2)
        assert top == 2.0, f"top_margin = {top} см, ожидалось 2.0"
        assert bottom == 2.0, f"bottom_margin = {bottom} см, ожидалось 2.0"
        assert left == 2.0, f"left_margin = {left} см, ожидалось 2.0"
        assert right == 1.5, f"right_margin = {right} см, ожидалось 1.5"


def test_kp_numbers_use_nbsp():
    """В ячейках Цена (tc[3]) и Сумма (tc[4]) разделитель тысяч —
    non-breaking space (U+00A0), а не обычный пробел. Это не даёт Word'у
    переносить «37 338» на две строки.
    """
    # 100 USD * 90 ₽ * 1.15 = 10 350; qty=4 → итог 41 400. В обоих
    # ячейках появятся пробелы как разделители тысяч.
    doc = _build_kp([(100.0, 4, "ЧислоNBSP")])
    # Этап 9А.2.7: внешняя обёрточная таблица убрана; таблица позиций —
    # единственная в body.
    inner = doc.tables[0]._tbl
    rows = inner.findall(f"{_NS}tr")
    # rows[0] — header, rows[1] — data, rows[-1] — total.
    data_tcs = rows[1].findall(f"{_NS}tc")
    for col_idx in (3, 4):  # Цена, Сумма
        text = "".join(
            t.text or "" for t in data_tcs[col_idx].findall(f".//{_NS}t")
        )
        assert " " in text, (
            f"В колонке {col_idx} ('{text}') нет non-breaking space — "
            "Word будет рвать число между разрядами"
        )
        assert " " not in text, (
            f"В колонке {col_idx} ('{text}') остался обычный пробел "
            "между разрядами — должен быть NBSP"
        )
    # Тоже для значения ИТОГО (последняя ячейка последней строки).
    itogo_text = "".join(
        t.text or ""
        for t in rows[-1].findall(f"{_NS}tc")[-1].findall(f".//{_NS}t")
    )
    assert " " in itogo_text and " " not in itogo_text, (
        f"ИТОГО '{itogo_text}': разделитель тысяч должен быть NBSP"
    )


def test_kp_table_centered():
    """Внешняя обёрточная таблица расположена симметрично относительно
    текстовой зоны: tblInd = 0 (или отсутствует) + tblJc = 'center'.
    """
    doc = _build_kp([(100.0, 1, "Центр")])
    outer = doc.tables[0]._tbl
    tblPr = outer.find(f"{_NS}tblPr")
    assert tblPr is not None

    # Отсутствие отрицательного tblInd (или вообще отсутствие).
    tblInd = tblPr.find(f"{_NS}tblInd")
    if tblInd is not None:
        val = int(tblInd.get(f"{_NS}w"))
        assert val == 0, (
            f"tblInd = {val}, ожидался 0 или отсутствие узла"
        )

    # tblJc центрирует таблицу относительно полей страницы.
    tblJc = tblPr.find(f"{_NS}jc")
    assert tblJc is not None, (
        "У внешней таблицы не задан <w:jc/> — она будет прижата к левому "
        "полю и выглядеть несимметрично"
    )
    assert tblJc.get(f"{_NS}val") == "center", (
        f"tblJc.val = {tblJc.get(f'{_NS}val')!r}, ожидался 'center'"
    )
