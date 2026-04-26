# Тесты KP-генератора (этап 8.2): kp_builder.build_kp_docx +
# роутер /project/{id}/export/kp.
#
# exchange_rate мокается, чтобы не ходить к ЦБ РФ; spec_items — либо
# создаются через реальный select_variant (интеграционные тесты), либо
# мокается list_spec_items (unit-тесты точного округления).

from __future__ import annotations

import json
from datetime import date
from decimal import Decimal
from io import BytesIO
from unittest.mock import patch

import pytest
from sqlalchemy import text as _t

import docx as _docx

from app.services import spec_service
from app.services.export import kp_builder


# ---------- helpers -------------------------------------------------------

def _seed_user(db_session, *, login: str = "kp-manager") -> int:
    from app.auth import hash_password
    row = db_session.execute(
        _t(
            "INSERT INTO users (login, password_hash, role, name) "
            "VALUES (:l, :p, 'manager', 'KP Manager') RETURNING id"
        ),
        {"l": login, "p": hash_password("x")},
    ).first()
    db_session.commit()
    return int(row.id)


def _insert_cpu(db_session, *, sku: str = "CPU-KP", gtin: str = "9901010101010") -> int:
    row = db_session.execute(
        _t(
            "INSERT INTO cpus "
            "  (sku, manufacturer, model, socket, cores, threads, "
            "   base_clock_ghz, turbo_clock_ghz, tdp_watts, gtin) "
            "VALUES "
            "  (:sku, 'Intel Corporation', 'Intel Core i5-12400F', 'LGA1700', "
            "   6, 12, 2.5, 4.4, 65, :gtin) RETURNING id"
        ),
        {"sku": sku, "gtin": gtin},
    ).first()
    db_session.commit()
    return int(row.id)


def _insert_ram(db_session, *, sku: str = "RAM-KP") -> int:
    row = db_session.execute(
        _t(
            "INSERT INTO rams "
            "  (sku, manufacturer, model, memory_type, form_factor, "
            "   module_size_gb, modules_count, frequency_mhz, gtin) "
            "VALUES "
            "  (:sku, 'Kingston', 'Kingston 8GB DDR4', 'DDR4', 'DIMM', "
            "   8, 2, 3200, '0202020202023') RETURNING id"
        ),
        {"sku": sku},
    ).first()
    db_session.commit()
    return int(row.id)


def _make_query(
    db_session,
    *,
    project_id: int,
    user_id: int,
    cpu_id: int,
    ram_id: int,
    total_usd: float = 220.0,
) -> int:
    build_result = {
        "status": "ok",
        "variants": [
            {
                "manufacturer": "Intel",
                "path_used":    "default",
                "used_transit": False,
                "total_usd":    total_usd,
                "total_rub":    total_usd * 90.0,
                "components": [
                    {
                        "category": "cpu", "component_id": cpu_id,
                        "model": "Intel Core i5-12400F", "sku": "BX8071512400F",
                        "manufacturer": "Intel", "quantity": 1,
                        "price_usd": 180.0, "price_rub": 16200.0,
                        "supplier": "OCS", "supplier_sku": "cpu-kp",
                    },
                    {
                        "category": "ram", "component_id": ram_id,
                        "model": "Kingston 8GB DDR4", "sku": "KVR-KP",
                        "manufacturer": "Kingston", "quantity": 2,
                        "price_usd": 20.0, "price_rub": 1800.0,
                        "supplier": "OCS", "supplier_sku": "ram-kp",
                    },
                ],
                "warnings": [],
            },
        ],
        "refusal_reason": None,
        "usd_rub_rate":   90.0,
        "fx_source":      "fallback",
    }
    row = db_session.execute(
        _t(
            "INSERT INTO queries "
            "  (project_id, user_id, raw_text, build_result_json, status, "
            "   cost_usd, cost_rub) "
            "VALUES (:pid, :uid, :rt, CAST(:br AS JSONB), 'ok', 0, 0) "
            "RETURNING id"
        ),
        {
            "pid": project_id, "uid": user_id, "rt": "тест KP",
            "br": json.dumps(build_result, ensure_ascii=False),
        },
    ).first()
    db_session.commit()
    return int(row.id)


def _mock_rate(value: str = "95", rate_date: date | None = None):
    return patch(
        "app.services.export.kp_builder.exchange_rate.get_usd_rate",
        return_value=(Decimal(value), rate_date or date(2026, 4, 24), "cache"),
    )


def _inner_table(doc):
    # Этап 9А.2.7: после программной сборки КП внешняя обёрточная
    # таблица убрана из шаблона; таблица позиций — единственная в body.
    return doc.tables[0]


def _data_rows_texts(doc):
    """Возвращает список списков tc.text для строк данных таблицы KP
    (кроме заголовка и строки ИТОГО). Этап 9А.2.5: разделитель тысяч
    в числах теперь — non-breaking space (U+00A0); чтобы существующие
    ассерты остались читаемыми с обычным пробелом, нормализуем NBSP→
    обычный пробел при чтении из документа."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    inner = _inner_table(doc)
    rows = inner._tbl.findall(f"{ns}tr")
    out = []
    for tr in rows[1:-1]:
        tcs = tr.findall(f"{ns}tc")
        texts = []
        for tc in tcs:
            ts = [t.text or "" for t in tc.findall(f".//{ns}t")]
            texts.append("".join(ts).replace(" ", " "))
        out.append(texts)
    return out


def _itogo_value(doc) -> str:
    """Возвращает текст последней ячейки строки ИТОГО (NBSP → пробел,
    см. комментарий выше)."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    inner = _inner_table(doc)
    rows = inner._tbl.findall(f"{ns}tr")
    itogo_tcs = rows[-1].findall(f"{ns}tc")
    ts = [t.text or "" for t in itogo_tcs[-1].findall(f".//{ns}t")]
    return "".join(ts).replace(" ", " ")


def _header_date(doc) -> str | None:
    """Возвращает содержимое параграфа «№ б/н от DD.MM.YYYYг.»."""
    for p in doc.paragraphs:
        if "№" in p.text and "от" in p.text:
            return p.text
    return None


# ---------- unit-тесты с моком list_spec_items ----------------------------

def _fake_items(items_spec: list[tuple[float, int, str]]) -> list[dict]:
    """[(unit_usd, qty, name), …] → spec-items dicts."""
    out = []
    for i, (uu, qty, name) in enumerate(items_spec, start=1):
        out.append({
            "id": i, "query_id": 100 + i, "variant_manufacturer": "Intel",
            "quantity": qty, "position": i,
            "auto_name": name, "custom_name": None, "display_name": name,
            "unit_usd": uu, "unit_rub": 0.0,
            "total_usd": round(uu * qty, 2), "total_rub": 0.0,
            "created_at": None, "updated_at": None,
        })
    return out


def test_rounding_math_matches_spec_example():
    """unit_usd=100.10, rate=95, markup=15 → sell = 10 937 (без «руб.»).

    Новая раскладка этапа 8.4: tc[2]=кол-во, tc[3]=цена, tc[4]=сумма.
    """
    items = _fake_items([(100.10, 1, "Конфигурация точного теста")])
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=items,
    ), _mock_rate("95"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=15, db=None,
        )
    doc = _docx.Document(BytesIO(data))
    rows = _data_rows_texts(doc)
    assert len(rows) == 1
    # tc[0]=№, tc[1]=имя, tc[2]=кол-во, tc[3]=цена, tc[4]=сумма
    assert rows[0][0] == "1"
    assert rows[0][1] == "Конфигурация точного теста"
    assert rows[0][2] == "1"
    assert rows[0][3] == "10 937"
    assert rows[0][4] == "10 937"
    assert _itogo_value(doc) == "10 937"


def test_markup_zero_means_no_markup():
    """markup=0 → продажа = math.ceil(base), без надбавки."""
    items = _fake_items([(100.0, 1, "No markup")])
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=items,
    ), _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=0, db=None,
        )
    rows = _data_rows_texts(_docx.Document(BytesIO(data)))
    # 100.0 * 90 = 9000 → +0% → 9000 (цена в tc[3], сумма в tc[4])
    assert rows[0][3] == "9 000"
    assert rows[0][4] == "9 000"


def test_markup_500_is_accepted_and_applied():
    """markup=500 — граничное значение, документ генерируется."""
    items = _fake_items([(1.0, 1, "Граничная")])
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=items,
    ), _mock_rate("100"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=500, db=None,
        )
    rows = _data_rows_texts(_docx.Document(BytesIO(data)))
    # 1 * 100 = 100 → +500% → 600 (цена — в tc[3])
    assert rows[0][3] == "600"


def test_markup_negative_raises_value_error():
    with pytest.raises(ValueError):
        kp_builder.build_kp_docx(project_id=1, markup_percent=-5, db=None)


def test_markup_501_raises_value_error():
    with pytest.raises(ValueError):
        kp_builder.build_kp_docx(project_id=1, markup_percent=501, db=None)


def test_two_configs_produce_two_numbered_rows():
    items = _fake_items([
        (100.0, 2, "Первая конфигурация"),
        (200.0, 1, "Вторая конфигурация"),
    ])
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=items,
    ), _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=10, db=None,
        )
    rows = _data_rows_texts(_docx.Document(BytesIO(data)))
    assert len(rows) == 2
    assert rows[0][0] == "1"
    assert rows[1][0] == "2"
    assert rows[0][1] == "Первая конфигурация"
    assert rows[1][1] == "Вторая конфигурация"
    # Новый порядок: tc[2]=кол-во, tc[3]=цена, tc[4]=сумма
    # 100*90=9000 → +10% = 9900, qty=2, итого строки 19800
    assert rows[0][2] == "2"
    assert rows[0][3] == "9 900"
    assert rows[0][4] == "19 800"
    # 200*90=18000 → +10% = 19800, qty=1
    assert rows[1][2] == "1"
    assert rows[1][3] == "19 800"
    assert rows[1][4] == "19 800"
    # ИТОГО = 19800 + 19800 = 39600
    assert _itogo_value(_docx.Document(BytesIO(data))) == "39 600"


def test_header_date_replaced_with_today():
    items = _fake_items([(100.0, 1, "Для даты")])
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=items,
    ), _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=0, db=None,
        )
    doc = _docx.Document(BytesIO(data))
    header = _header_date(doc)
    assert header is not None
    today_str = date.today().strftime("%d.%m.%Y")
    assert today_str in header, f"Сегодняшняя дата не подставлена: {header!r}"


# ---------- интеграционный тест через реальную БД -------------------------

def test_single_config_integration(db_session):
    """Генерация для проекта с 1 конфигурацией через реальную select_variant."""
    uid = _seed_user(db_session)
    pid = spec_service.create_empty_project(
        db_session, user_id=uid, name="KP интеграция",
    )
    cpu_id = _insert_cpu(db_session)
    ram_id = _insert_ram(db_session)
    qid = _make_query(
        db_session, project_id=pid, user_id=uid,
        cpu_id=cpu_id, ram_id=ram_id, total_usd=220.0,
    )
    spec_service.select_variant(
        db_session, project_id=pid, query_id=qid,
        manufacturer="Intel", quantity=1,
    )

    with _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=pid, markup_percent=15, db=db_session,
        )
    doc = _docx.Document(BytesIO(data))
    rows = _data_rows_texts(doc)
    assert len(rows) == 1
    assert rows[0][0] == "1"
    # 220 * 90 = 19800 → +15% → math.ceil(22770) = 22770
    # tc[2]=кол-во, tc[3]=цена, tc[4]=сумма (новый порядок)
    assert rows[0][2] == "1"
    assert rows[0][3] == "22 770"
    assert rows[0][4] == "22 770"
    assert "Системный блок" in rows[0][1]


# ---------- тесты роутера /export/kp --------------------------------------

def _router_mock_rate():
    """patch exchange_rate.get_usd_rate, используется kp_builder внутри
    обработчика /export/kp."""
    return _mock_rate("90")


def _make_project_with_spec(db_session, user_id: int, project_name: str):
    pid = spec_service.create_empty_project(
        db_session, user_id=user_id, name=project_name,
    )
    cpu_id = _insert_cpu(db_session, sku=f"CPU-{project_name}",
                         gtin=f"900{user_id:04d}{pid:06d}")
    ram_id = _insert_ram(db_session, sku=f"RAM-{project_name}")
    qid = _make_query(
        db_session, project_id=pid, user_id=user_id,
        cpu_id=cpu_id, ram_id=ram_id, total_usd=220.0,
    )
    spec_service.select_variant(
        db_session, project_id=pid, query_id=qid,
        manufacturer="Intel", quantity=1,
    )
    return pid


def test_kp_endpoint_returns_docx(
    db_session, manager_client, manager_user,
):
    pid = _make_project_with_spec(db_session, manager_user["id"], "Эндпоинт КП")
    with _router_mock_rate():
        r = manager_client.get(f"/project/{pid}/export/kp?markup=15")
    assert r.status_code == 200, r.text[:300]
    assert "wordprocessingml" in r.headers["content-type"]
    cd = r.headers["content-disposition"]
    assert cd.startswith("attachment;")
    assert "filename*=UTF-8''" in cd
    # Бинарь валиден — открывается python-docx.
    doc = _docx.Document(BytesIO(r.content))
    assert doc.tables  # хотя бы одна таблица


def test_kp_endpoint_default_markup_is_15(
    db_session, manager_client, manager_user,
):
    """markup по умолчанию = 15% — такое же поведение, как с явным 15."""
    pid = _make_project_with_spec(db_session, manager_user["id"], "default-markup")
    with _router_mock_rate():
        r_default = manager_client.get(f"/project/{pid}/export/kp")
        r_explicit = manager_client.get(f"/project/{pid}/export/kp?markup=15")
    assert r_default.status_code == 200
    assert r_explicit.status_code == 200
    # Сравниваем цены в сгенерированных docx. Разные бинари
    # (разные rsid/временные метки внутри), но цены обязаны совпадать.
    rows_a = _data_rows_texts(_docx.Document(BytesIO(r_default.content)))
    rows_b = _data_rows_texts(_docx.Document(BytesIO(r_explicit.content)))
    # После этапа 8.4 цена переехала в tc[3], сумма — в tc[4].
    assert rows_a[0][3] == rows_b[0][3]
    assert rows_a[0][4] == rows_b[0][4]


def test_kp_endpoint_rejects_markup_above_500(
    db_session, manager_client, manager_user,
):
    pid = _make_project_with_spec(db_session, manager_user["id"], "too-much")
    with _router_mock_rate():
        r = manager_client.get(f"/project/{pid}/export/kp?markup=600")
    assert r.status_code == 400


def test_kp_endpoint_forbidden_for_other_user(
    db_session, app_client, manager_user, manager2_user,
):
    pid = spec_service.create_empty_project(
        db_session, user_id=manager_user["id"], name="Чужой KP",
    )
    from tests.test_web.conftest import _login
    _login(app_client, manager2_user["login"], manager2_user["password"])
    with _router_mock_rate():
        r = app_client.get(f"/project/{pid}/export/kp?markup=15")
    assert r.status_code == 403


def test_kp_endpoint_404_for_missing_project(manager_client):
    with _router_mock_rate():
        r = manager_client.get("/project/999999/export/kp?markup=15")
    assert r.status_code == 404


# ---------- этап 9А.2.7: программная сборка таблицы -----------------------

_NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _build_doc(items, *, markup=15, rate="95"):
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=items,
    ), _mock_rate(rate):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=markup, db=None,
        )
    return _docx.Document(BytesIO(data))


def test_kp_builds_5_columns_correctly():
    """После генерации — 5 колонок с заданными ширинами в tblGrid."""
    doc = _build_doc(_fake_items([(100.0, 1, "Конф")]))
    inner = _inner_table(doc)
    grid = inner._tbl.find(f"{_NS_W}tblGrid")
    cols = grid.findall(f"{_NS_W}gridCol")
    assert len(cols) == 5
    widths = [int(c.get(f"{_NS_W}w")) for c in cols]
    assert widths == [454, 3686, 794, 2041, 2098]


def test_kp_header_row_text():
    """Шапка содержит ровно 5 заголовков из спецификации."""
    doc = _build_doc(_fake_items([(100.0, 1, "X")]))
    inner = _inner_table(doc)
    rows = inner._tbl.findall(f"{_NS_W}tr")
    header_tcs = rows[0].findall(f"{_NS_W}tc")
    titles = []
    for tc in header_tcs:
        ts = [t.text or "" for t in tc.findall(f".//{_NS_W}t")]
        titles.append("".join(ts))
    assert titles == [
        "№ п/п", "Наименование", "Кол-во",
        "Цена с НДС (руб.)", "Сумма с НДС (руб.)",
    ]


def test_kp_data_row_complete():
    """Каждая data-строка имеет все 5 заполненных ячеек."""
    items = _fake_items([
        (100.0, 2, "Конфигурация А"),
        (200.0, 1, "Конфигурация Б"),
    ])
    doc = _build_doc(items, markup=10, rate="90")
    rows = _data_rows_texts(doc)
    assert len(rows) == 2
    for row in rows:
        assert len(row) == 5
        for cell in row:
            assert cell != "", f"Пустая ячейка: {row}"


def test_kp_uses_calibri_font():
    """В каждом run таблицы шрифт явно = Calibri (rPr/rFonts)."""
    doc = _build_doc(_fake_items([(100.0, 1, "Конф")]))
    inner = _inner_table(doc)
    runs = inner._tbl.findall(f".//{_NS_W}r")
    assert runs, "В таблице нет runs"
    for r in runs:
        rfonts = r.find(f"{_NS_W}rPr/{_NS_W}rFonts")
        assert rfonts is not None, "У run нет rPr/rFonts — шрифт унаследуется"
        for ax in ("ascii", "hAnsi", "cs", "eastAsia"):
            val = rfonts.get(f"{_NS_W}{ax}")
            assert val == "Calibri", f"run uses {ax}={val!r}, expected Calibri"


def test_kp_no_inherited_times_new_roman():
    """Ни в одной ячейке таблицы не наследуется TNR из Normal-стиля."""
    doc = _build_doc(_fake_items([(100.0, 1, "Конф")]))
    inner = _inner_table(doc)
    for rfonts in inner._tbl.findall(f".//{_NS_W}rFonts"):
        for ax in ("ascii", "hAnsi", "cs", "eastAsia"):
            val = rfonts.get(f"{_NS_W}{ax}")
            if val:
                assert val != "Times New Roman", \
                    f"в таблице остался TNR на оси {ax}"


def test_kp_itogo_gridspan():
    """Последняя строка имеет 2 видимые ячейки; первая с gridSpan=4."""
    doc = _build_doc(_fake_items([(100.0, 1, "X")]))
    inner = _inner_table(doc)
    rows = inner._tbl.findall(f"{_NS_W}tr")
    itogo_tcs = rows[-1].findall(f"{_NS_W}tc")
    assert len(itogo_tcs) == 2
    gs = itogo_tcs[0].find(f"{_NS_W}tcPr/{_NS_W}gridSpan")
    assert gs is not None
    assert gs.get(f"{_NS_W}val") == "4"


def test_kp_table_below_kp_title():
    """Порядок body: реквизиты → дата → «Коммерческое предложение»
    → таблица → подпись/печать."""
    doc = _build_doc(_fake_items([(100.0, 1, "Конф")]))
    body = doc.element.body
    seq = []
    for ch in list(body):
        tag = ch.tag.split("}")[-1]
        if tag == "p":
            text = "".join((t.text or "") for t in ch.iter(f"{_NS_W}t"))
            has_inline = ch.find(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
            ) is not None
            if "quadro.tatar" in text.lower():
                seq.append("requisites_end")
            elif text.startswith("№ б/н"):
                seq.append("date")
            elif "Коммерческое предложение" in text:
                seq.append("kp_title")
            elif has_inline:
                seq.append("signature")
        elif tag == "tbl":
            seq.append("table")
    # Каждый ключевой элемент должен быть на своём месте и в этом порядке.
    for label in ("requisites_end", "date", "kp_title", "table", "signature"):
        assert label in seq, f"В порядке body нет {label}: {seq}"
    indexes = [seq.index(l) for l in
               ("requisites_end", "date", "kp_title", "table", "signature")]
    assert indexes == sorted(indexes), \
        f"Порядок элементов нарушен: {seq}"


def test_kp_kp_title_is_centered_bold_14pt():
    """Заголовок «Коммерческое предложение» — center, bold, 14pt (sz=28)."""
    doc = _build_doc(_fake_items([(100.0, 1, "X")]))
    body = doc.element.body
    for p in body.findall(f"{_NS_W}p"):
        text = "".join((t.text or "") for t in p.iter(f"{_NS_W}t"))
        if "Коммерческое предложение" in text:
            jc = p.find(f"{_NS_W}pPr/{_NS_W}jc")
            assert jc is not None and jc.get(f"{_NS_W}val") == "center"
            r = p.find(f"{_NS_W}r")
            assert r.find(f"{_NS_W}rPr/{_NS_W}b") is not None
            sz = r.find(f"{_NS_W}rPr/{_NS_W}sz")
            assert sz is not None and sz.get(f"{_NS_W}val") == "28"
            return
    pytest.fail("Заголовок «Коммерческое предложение» не найден")


def test_kp_signature_image_preserved():
    """Inline-картинка подписи/печати сохранена в body."""
    doc = _build_doc(_fake_items([(100.0, 1, "X")]))
    body = doc.element.body
    inline_pics = body.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
    )
    assert len(inline_pics) >= 1, "Inline-картинка подписи не сохранилась"


def test_kp_normal_style_uses_calibri():
    """Normal-стиль шаблона переведён на Calibri 11pt (а не TNR 14pt)."""
    import zipfile
    from pathlib import Path
    template = Path(kp_builder._TEMPLATE_PATH)
    with zipfile.ZipFile(template) as z:
        styles_xml = z.read("word/styles.xml").decode("utf-8")
    import re
    m = re.search(
        r'<w:style[^>]*w:styleId="a"[^>]*>.*?</w:style>',
        styles_xml, re.DOTALL,
    )
    assert m, "Normal style 'a' не найден в styles.xml"
    body = m.group(0)
    rfonts_m = re.search(r'<w:rFonts([^/]*)/?>', body)
    assert rfonts_m, "В Normal стиле нет rFonts"
    rfonts_attrs = rfonts_m.group(1)
    assert 'Calibri' in rfonts_attrs
    assert 'Times New Roman' not in rfonts_attrs
    sz_m = re.search(r'<w:sz w:val="(\d+)"/>', body)
    assert sz_m and sz_m.group(1) == "22", \
        f"Normal sz={sz_m.group(1) if sz_m else 'None'}, ожидалось 22 (11pt)"
