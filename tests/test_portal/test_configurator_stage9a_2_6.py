"""Тесты этапа 9А.2.6.

Покрывают добивающие фиксы Word KP:
  A. Расширенные колонки Цена (3.6 см) / Сумма (3.7 см) и сокращённое
     «Наименование» (6.5 см). На предыдущей раскладке (2.7/2.8 см) Word
     рвал «40 156» на «40 15 / 6» из-за padding'а ячейки, даже при
     заданном <w:noWrap/>.
  B. Удаление хвостового пустого параграфа с _GoBack-bookmark
     (визуально отображался как «□» в правом нижнем углу под печатью).
  C. Корректный пробел перед датой в строке «№ б/н от <дата>г.» —
     перезаписываем параграф одним run'ом с xml:space="preserve".
"""

from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO
from unittest.mock import patch

import docx as _docx

from portal.services.configurator.export import kp_builder

_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_XMLNS = "{http://www.w3.org/XML/1998/namespace}"


def _fake_items(items_spec):
    out = []
    for i, (uu, qty, name) in enumerate(items_spec, start=1):
        out.append({
            "id": i, "query_id": 100 + i, "variant_manufacturer": "Intel",
            "quantity": qty, "position": i,
            "auto_name": name, "custom_name": None, "display_name": name,
            "unit_usd": uu, "unit_rub": 0.0,
            "total_usd": round(uu * qty, 2), "total_rub": 0.0,
            "created_at": None, "updated_at": None,
        })
    return out


def _mock_rate(value: str = "90"):
    return patch(
        "portal.services.configurator.export.kp_builder.exchange_rate.get_usd_rate",
        return_value=(Decimal(value), date(2026, 4, 25), "cache"),
    )


def _build_kp(items_spec, markup=15):
    with patch(
        "portal.services.configurator.export.kp_builder.spec_service.list_spec_items",
        return_value=_fake_items(items_spec),
    ), _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=markup, db=None,
        )
    return _docx.Document(BytesIO(data))


# =====================================================================
# A. Ширины колонок: Цена/Сумма расширены, Наименование сокращено
# =====================================================================

def test_kp_column_widths_v2():
    """gridCol и tcW во ВСЕХ строках имеют новые ширины:
    454 / 3686 / 794 / 2041 / 2098 (= 9073 twips).

    Старые значения 4706/1531/1588 на печати давали разрыв чисел
    «40 156» на «40 15 / 6».
    """
    expected = (454, 3686, 794, 2041, 2098)
    assert sum(expected) == 9073

    doc = _build_kp([(2_222.0, 7, "Дорогой блок ради больших чисел")])
    # Этап 9А.2.7: внешняя обёрточная таблица убрана; таблица позиций —
    # единственная в body.
    inner_tbl = doc.tables[0]._tbl

    # tblGrid
    grid = inner_tbl.find(f"{_NS}tblGrid")
    assert grid is not None
    grid_widths = tuple(
        int(gc.get(f"{_NS}w"))
        for gc in grid.findall(f"{_NS}gridCol")
    )
    assert grid_widths == expected, (
        f"tblGrid widths = {grid_widths}, ожидались {expected}"
    )

    # tcW в каждой строке (header + data + total). У строки ИТОГО
    # 4 первые ячейки объединены через gridSpan=4, поэтому tcW первой
    # ячейки = сумма первых четырёх gridCol; tcW второй = последний.
    rows = inner_tbl.findall(f"{_NS}tr")
    header = rows[0]
    header_widths = tuple(
        int(tc.find(f"{_NS}tcPr").find(f"{_NS}tcW").get(f"{_NS}w"))
        for tc in header.findall(f"{_NS}tc")
    )
    assert header_widths == expected, (
        f"header tcW = {header_widths}, ожидались {expected}"
    )

    # data-строка: ровно ту же раскладку
    data_widths = tuple(
        int(tc.find(f"{_NS}tcPr").find(f"{_NS}tcW").get(f"{_NS}w"))
        for tc in rows[1].findall(f"{_NS}tc")
    )
    assert data_widths == expected, (
        f"data tcW = {data_widths}, ожидались {expected}"
    )

    # ИТОГО: span 4 + последняя
    total_tcs = rows[-1].findall(f"{_NS}tc")
    total_widths = tuple(
        int(tc.find(f"{_NS}tcPr").find(f"{_NS}tcW").get(f"{_NS}w"))
        for tc in total_tcs
    )
    assert total_widths == (sum(expected[:4]), expected[4]), (
        f"total tcW = {total_widths}, ожидались "
        f"({sum(expected[:4])}, {expected[4]})"
    )


def test_kp_price_column_wider_than_old():
    """Sanity-check: ширина Цена ≥ 2000 twips (≥ 3.5 см). Старая
    раскладка давала 1531 — на ней шестизначные числа рвались."""
    doc = _build_kp([(100.0, 1, "Sanity")])
    # Этап 9А.2.7: внешняя обёрточная таблица убрана; таблица позиций —
    # единственная в body.
    inner_tbl = doc.tables[0]._tbl
    grid_widths = [
        int(gc.get(f"{_NS}w"))
        for gc in inner_tbl.find(f"{_NS}tblGrid").findall(f"{_NS}gridCol")
    ]
    assert grid_widths[3] >= 2000, (
        f"Цена ширина = {grid_widths[3]} twips, должно быть ≥ 2000"
    )
    assert grid_widths[4] >= 2000, (
        f"Сумма ширина = {grid_widths[4]} twips, должно быть ≥ 2000"
    )


# =====================================================================
# B. Хвостовой пустой параграф удалён
# =====================================================================

def test_kp_no_trailing_empty_paragraph():
    """В body документа НЕ должно быть пустого <w:p/> непосредственно
    перед финальным <w:sectPr>. Исходный kp_template.docx содержит
    такой параграф с <w:bookmarkStart name="_GoBack"/> и пустым
    содержимым, который Word рендерит как «□» в правом нижнем углу
    страницы под печатью.
    """
    doc = _build_kp([(100.0, 1, "Бэйкендный текст")])
    body = doc.element.body
    children = list(body)
    assert children, "body пустой — что-то сломалось"
    # Последний элемент — sectPr.
    assert children[-1].tag == f"{_NS}sectPr", (
        f"Последний элемент body = {children[-1].tag}, ожидался sectPr"
    )
    # Перед sectPr — НЕ пустой параграф. Допускается:
    #   а) tbl (наша КП-таблица)
    #   б) p со значимым контентом
    pre_sectpr = children[-2]
    if pre_sectpr.tag == f"{_NS}p":
        # Проверяем, что параграф НЕ пустой (есть text run или drawing).
        has_meaningful_content = False
        for r in pre_sectpr.findall(f"{_NS}r"):
            for t in r.findall(f"{_NS}t"):
                if (t.text or "").strip():
                    has_meaningful_content = True
                    break
            if has_meaningful_content:
                break
            if (r.find(f".//{_NS}drawing") is not None
                    or r.find(f".//{_NS}pict") is not None):
                has_meaningful_content = True
                break
        assert has_meaningful_content, (
            "Перед <w:sectPr> остался пустой <w:p/> "
            "(скорее всего хвост с _GoBack-bookmark) — "
            "это и есть «□» под печатью"
        )


def test_kp_goback_bookmark_removed():
    """В готовом документе НЕТ <w:bookmarkStart name="_GoBack">
    (Word'овская точка «последнее редактирование», ездит вместе с
    хвостовым пустым параграфом и не несёт смысловой нагрузки в
    автогенерируемом КП).
    """
    doc = _build_kp([(100.0, 1, "Без GoBack")])
    body = doc.element.body
    for bm in body.findall(f".//{_NS}bookmarkStart"):
        name = bm.get(f"{_NS}name") or ""
        assert name != "_GoBack", (
            "В готовом документе остался <w:bookmarkStart name=\"_GoBack\">,"
            " хотя именно этот пустой параграф должен был быть удалён"
        )


# =====================================================================
# C. Дата в шапке: «№ б/н от <space><дата>г.»
# =====================================================================

def _header_paragraph(doc):
    """Параграф с «№ б/н от …»."""
    for p in doc.paragraphs:
        if "№" in p.text and "от" in p.text:
            return p
    return None


def test_kp_date_has_space_before_date():
    """Финальный текст параграфа: «№ б/н от 25.04.2026г.» — между
    «от» и датой РОВНО один пробел, без слипания «от25.04.2026»
    и без артефактов от старых proofErr-runs.
    """
    today = date.today().strftime("%d.%m.%Y")
    doc = _build_kp([(100.0, 1, "Дата")])
    p = _header_paragraph(doc)
    assert p is not None, "В документе нет параграфа с № б/н"
    text = p.text
    # Проверяем итоговую строку буквально.
    expected = f"№ б/н от {today}г."
    assert text == expected, (
        f"Текст параграфа = {text!r}, ожидался {expected!r}"
    )
    # Дополнительная страховка: нет слипания «от<число>»
    assert f"от{today}" not in text, (
        f"В тексте {text!r} «от» слиплось с датой"
    )


def test_kp_date_paragraph_uses_xml_space_preserve():
    """Параграф «№ б/н от …» содержит ровно один <w:r>, и его <w:t>
    несёт xml:space="preserve". Это гарантирует, что Word не схлопнет
    пробел перед датой даже при перерисовке.
    """
    doc = _build_kp([(100.0, 1, "preserve")])
    p = _header_paragraph(doc)
    assert p is not None
    runs = p._p.findall(f"{_NS}r")
    assert len(runs) == 1, (
        f"Ожидался ровно один run в параграфе с датой, нашлось {len(runs)}. "
        "Старые proofErr/run-фрагменты должны быть полностью вычищены."
    )
    ts = runs[0].findall(f"{_NS}t")
    assert len(ts) == 1
    assert ts[0].get(f"{_XMLNS}space") == "preserve", (
        "У <w:t> с датой нет xml:space=\"preserve\" — "
        "Word может схлопнуть пробел перед датой"
    )


def test_kp_date_paragraph_no_proof_err_remnants():
    """В параграфе с датой НЕ должно остаться <w:proofErr/> элементов
    (gramStart/gramEnd) из исходного шаблона. На прошлой реализации они
    оставались висеть после перезаписи runs[0] и могли визуально
    проявляться как лишние glyph'ы.
    """
    doc = _build_kp([(100.0, 1, "no-proofErr")])
    p = _header_paragraph(doc)
    assert p is not None
    proof_errs = p._p.findall(f"{_NS}proofErr")
    assert proof_errs == [], (
        f"В параграфе с датой остались proofErr-узлы: {proof_errs}"
    )
