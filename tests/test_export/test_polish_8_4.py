# Юнит-тесты правок этапа 8.4:
# - санитизация имён файлов экспорта (Excel/КП);
# - структура Excel после удаления «Создан: …»: нет A2, формулы в N,
#   ИТОГО суммирует только строки системных блоков;
# - KP: новый порядок колонок, отсутствие «руб.» в значениях,
#   непустое ИТОГО с жирным форматированием;
# - GPU naming: «GeForce 210» вместо «RTX 1», fallback-подстраховка
#   на подозрительных моделях.

from __future__ import annotations

from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from unittest.mock import patch

import docx as _docx
from openpyxl import load_workbook

from app.routers.export_router import _safe_project_filename
from app.services.export import excel_builder, kp_builder
from app.services.spec_naming import (
    _block_gpu,
    _looks_suspicious_gpu_model,
    _short_gpu_model,
)


# =============================================================================
# Блок B — санитизация имён файлов
# =============================================================================

def test_safe_project_filename_replaces_windows_forbidden_chars():
    assert _safe_project_filename("Запрос от 24.04.2026 23:54", 1) == (
        "Запрос от 24.04.2026 23_54"
    )


def test_safe_project_filename_replaces_slashes():
    assert _safe_project_filename("Мой/проект*", 7) == "Мой_проект_"


def test_safe_project_filename_strips_trailing_space_and_dot():
    assert _safe_project_filename("Проект. ", 5) == "Проект"


def test_safe_project_filename_fallback_on_empty():
    assert _safe_project_filename("   ", 42) == "project_42"
    assert _safe_project_filename(None, 9) == "project_9"


def test_safe_project_filename_length_capped():
    long = "A" * 300
    result = _safe_project_filename(long, 1)
    assert len(result) <= 150
    assert result == "A" * 150


def test_safe_project_filename_all_forbidden_charset():
    # <>:"/\|?* — девять запрещённых Windows-символов подряд
    assert _safe_project_filename('a<>:"/\\|?*b', 1) == "a_________b"


# =============================================================================
# Блок C — Excel-раскладка: нет A2, формулы SUM в N на comp-строке
# =============================================================================

def _fake_blocks(item_name: str = "Системный блок") -> list:
    item = {
        "id": 1, "query_id": 101, "variant_manufacturer": "Intel",
        "quantity": 1, "position": 1,
        "auto_name": item_name, "custom_name": None,
        "display_name": item_name,
        "unit_usd": 200.0, "unit_rub": 18000.0,
        "total_usd": 200.0, "total_rub": 18000.0,
    }
    variant = {"manufacturer": "Intel"}
    comps = [
        {
            "category": "cpu", "component_id": 999,
            "model": "Intel Core i5", "sku": "CPU", "manufacturer": "Intel",
            "quantity": 1, "price_usd": 180.0,
            "specs_short": "6C/12T",
        },
        {
            "category": "ram", "component_id": 888,
            "model": "Kingston 8GB DDR4", "sku": "RAM", "manufacturer": "Kingston",
            "quantity": 2, "price_usd": 10.0,
            "specs_short": "8GB DDR4",
        },
    ]
    return [(item, variant, comps)]


def _build_xlsx_with_fake_data(item_name: str = "Системный блок"):
    """Собирает xlsx без обращения к БД: мокаем _load_project,
    spec_service.list_spec_items, _collect_blocks и _fetch_gtin_map.
    """
    fake_project = {
        "id": 1, "name": "Запрос от 24.04.2026 23:54",
        "created_at": datetime(2026, 4, 24, 23, 54),
        "author_login": "manager1", "author_name": "Менеджер",
    }
    blocks = _fake_blocks(item_name)

    with patch.object(excel_builder, "_load_project", return_value=fake_project), \
         patch(
             "app.services.export.excel_builder.spec_service.list_spec_items",
             return_value=[blocks[0][0]],
         ), \
         patch.object(excel_builder, "_collect_blocks", return_value=blocks), \
         patch.object(excel_builder, "_fetch_gtin_map", return_value={}):
        xlsx = excel_builder.build_project_xlsx(
            project_id=1, db=None,
            rate=Decimal("92.5"), rate_date=date(2026, 4, 24),
        )
    return load_workbook(BytesIO(xlsx))


def test_excel_no_cell_a2_after_polish():
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    # A2 пустая (нет «Создан: …»), но строка существует — иначе
    # N2/O2 курса съезжают (см. этап 8.6).
    a2 = ws["A2"].value
    assert not (a2 and "Создан" in str(a2)), (
        f"A2 не должен содержать «Создан»: {a2!r}"
    )


def test_excel_headers_in_row_3():
    """Этап 8.6: заголовки таблицы — в строке 3 (после шапки и пустой
    компактной строки 2 для ячеек курса)."""
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    assert (ws["D3"].value or "").startswith("Наименование") or \
           "аим" in (ws["D3"].value or "")


def test_excel_n_formula_on_comp_row_is_sum_of_components():
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    # Раскладка 8.6: row 3 — заголовки, row 4 — comp, rows 5-6 — компоненты.
    n_comp = ws["N4"].value
    assert isinstance(n_comp, str) and n_comp.startswith("=SUM(N"), (
        f"В N4 должна быть формула SUM по N компонентов, получили {n_comp!r}"
    )
    assert "N5" in n_comp and "N6" in n_comp
    assert ws["N5"].value == 180.0
    assert ws["N6"].value == 10.0


def test_excel_totals_sum_only_comp_rows():
    # Две конфигурации — должно быть =SUM(J<comp1>, J<comp2>),
    # не включая строки компонентов.
    item1 = {
        "id": 1, "query_id": 101, "variant_manufacturer": "Intel",
        "quantity": 1, "position": 1,
        "auto_name": "comp A", "custom_name": None, "display_name": "comp A",
        "unit_usd": 100.0, "unit_rub": 9000.0,
        "total_usd": 100.0, "total_rub": 9000.0,
    }
    item2 = {**item1, "id": 2, "query_id": 102, "auto_name": "comp B",
             "display_name": "comp B", "position": 2}
    blocks = [
        (item1, {"manufacturer": "Intel"}, [
            {"category": "cpu", "component_id": 1, "model": "CPU",
             "sku": "s1", "manufacturer": "Intel", "quantity": 1,
             "price_usd": 80.0, "specs_short": ""},
        ]),
        (item2, {"manufacturer": "AMD"}, [
            {"category": "cpu", "component_id": 2, "model": "CPU",
             "sku": "s2", "manufacturer": "AMD", "quantity": 1,
             "price_usd": 90.0, "specs_short": ""},
        ]),
    ]
    fake_project = {
        "id": 1, "name": "Два проекта",
        "created_at": datetime(2026, 4, 24, 12, 0),
        "author_login": "m", "author_name": "M",
    }
    with patch.object(excel_builder, "_load_project", return_value=fake_project), \
         patch(
             "app.services.export.excel_builder.spec_service.list_spec_items",
             return_value=[item1, item2],
         ), \
         patch.object(excel_builder, "_collect_blocks", return_value=blocks), \
         patch.object(excel_builder, "_fetch_gtin_map", return_value={}):
        xlsx = excel_builder.build_project_xlsx(
            project_id=1, db=None,
            rate=Decimal("90"), rate_date=date(2026, 4, 24),
        )
    wb = load_workbook(BytesIO(xlsx))
    ws = wb.active
    # Раскладка 8.6: row 3 заголовки, row 4 comp A, row 5 CPU A,
    # row 6 comp B, row 7 CPU B. last_data_row=7 → sum_row=9.
    g_formula = ws["G9"].value
    j_formula = ws["J9"].value
    assert g_formula == "=SUM(G4,G6)", f"Got {g_formula!r}"
    assert j_formula == "=SUM(J4,J6)", f"Got {j_formula!r}"


def test_excel_non_comp_rows_have_no_fill():
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    # Компонентные строки (5, 6) — без заливки.
    for r in (5, 6):
        for col in "BCDEFGHIJKLMN":
            cell = ws[f"{col}{r}"]
            fg = cell.fill.fgColor
            rgb = getattr(fg, "rgb", None) if fg else None
            if isinstance(rgb, str) and rgb != "00000000":
                # Допустимо только в comp-строке (row 4) — на ней голубая.
                assert False, (
                    f"Неожиданная заливка в {col}{r}: {rgb!r}"
                )


def test_excel_totals_cells_have_borders():
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    # Раскладка 8.6: last_data_row=6 → sum_row=8, pct_row=10, abs_row=11.
    for coord in ("F8", "G8", "I8", "J8", "I10", "J10", "I11", "J11"):
        b = ws[coord].border
        assert b.left and b.left.style, f"{coord}: нет левой границы"
        assert b.right and b.right.style, f"{coord}: нет правой границы"


# =============================================================================
# Этап 8.6 — Excel: курс, видимые границы, заголовки не на строке курса
# =============================================================================

def test_excel_kurs_position():
    """N1/O1 — «Курс на» + дата; N2/O2 — «Курс ЦБ» + число."""
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    n1 = (ws["N1"].value or "").strip()
    o1 = ws["O1"].value
    n2 = (ws["N2"].value or "").strip()
    o2 = ws["O2"].value
    assert "Курс" in n1, f"N1 должен содержать 'Курс …', получили {n1!r}"
    assert isinstance(o1, str) and "2026" in o1, f"O1 должна быть датой, {o1!r}"
    assert "Курс ЦБ" in n2, f"N2 должен быть 'Курс ЦБ', {n2!r}"
    assert isinstance(o2, (int, float)) and float(o2) > 0, (
        f"O2 должна быть числом курса, {o2!r}"
    )


def test_excel_no_kurs_in_header_row():
    """Заголовки таблицы (row 3) не должны иметь курс справа.

    После 8.6 курс лежит в N1/O1/N2/O2 — выше заголовков. Соответственно
    в N3/O3 должны быть либо None, либо текстовые заголовки колонок,
    но НЕ числовое значение курса и НЕ дата курса.
    """
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    n3 = ws["N3"].value
    o3 = ws["O3"].value
    # Числового значения курса в строке заголовков быть не должно.
    assert not isinstance(n3, (int, float)), f"N3 не должна быть числом, {n3!r}"
    assert not isinstance(o3, (int, float)), f"O3 не должна быть числом, {o3!r}"
    # Дата вида «дд.мм.гггг» в строке заголовков быть не должна.
    for v in (n3, o3):
        if isinstance(v, str):
            assert "2026" not in v, (
                f"В заголовках не должно быть даты курса: {v!r}"
            )


def test_excel_borders_visible_thin_black():
    """Этап 8.6: видимые границы — thin + чёрный по всем ячейкам таблицы."""
    wb = _build_xlsx_with_fake_data()
    ws = wb.active
    # Покрываем заголовки + данные (первые 3 строки данных).
    for row in (3, 4, 5, 6):
        for col in "ABCDEFGHIJKLMNO":
            cell = ws[f"{col}{row}"]
            for side_name in ("left", "right", "top", "bottom"):
                side = getattr(cell.border, side_name)
                assert side and side.style == "thin", (
                    f"{col}{row} side={side_name}: ожидался thin, получили {side.style!r}"
                )
                # Цвет — чёрный (FF000000) либо в формате 000000.
                rgb = getattr(side.color, "rgb", None) if side.color else None
                # Пустой rgb недопустим — проверяем только если задан.
                if rgb:
                    assert rgb.upper().endswith("000000"), (
                        f"{col}{row} side={side_name}: ожидался чёрный, {rgb!r}"
                    )


# =============================================================================
# Блок E — KP: порядок колонок, нет «руб.», ИТОГО заполнено и жирное
# =============================================================================

_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _kp_inner_table(doc):
    return doc.tables[0].rows[0].cells[0].tables[0]


def _kp_header_texts(doc) -> list[str]:
    inner = _kp_inner_table(doc)
    header_tr = inner._tbl.findall(f"{_NS}tr")[0]
    out = []
    for tc in header_tr.findall(f"{_NS}tc"):
        out.append("".join(t.text or "" for t in tc.findall(f".//{_NS}t")))
    return out


def _mock_rate(value: str = "90"):
    return patch(
        "app.services.export.kp_builder.exchange_rate.get_usd_rate",
        return_value=(Decimal(value), date(2026, 4, 24), "cache"),
    )


def _fake_items(items_spec):
    out = []
    for i, (uu, qty, name) in enumerate(items_spec, start=1):
        out.append({
            "id": i, "query_id": 100 + i, "variant_manufacturer": "Intel",
            "quantity": qty, "position": i,
            "auto_name": name, "custom_name": None, "display_name": name,
            "unit_usd": uu, "unit_rub": 0.0,
            "total_usd": round(uu * qty, 2), "total_rub": 0.0,
            "created_at": None, "updated_at": None,
        })
    return out


def _build_kp_with_fake(items_spec, markup=15):
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=_fake_items(items_spec),
    ), _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=markup, db=None,
        )
    return _docx.Document(BytesIO(data))


def test_kp_columns_order_kol_before_price():
    doc = _build_kp_with_fake([(100.0, 1, "Конфиг")])
    headers = _kp_header_texts(doc)
    # tc[0]=№, tc[1]=Наименование, tc[2]=Кол-во, tc[3]=Цена..., tc[4]=Сумма...
    assert len(headers) == 5
    assert "Кол" in headers[2], f"В колонке 2 должно быть «Кол-во»: {headers[2]!r}"
    assert "Цена" in headers[3], f"В колонке 3 — «Цена»: {headers[3]!r}"
    assert "Сумма" in headers[4], f"В колонке 4 — «Сумма»: {headers[4]!r}"


def test_kp_no_ruble_suffix_in_value_cells():
    doc = _build_kp_with_fake([(100.0, 1, "x")])
    inner = _kp_inner_table(doc)
    rows = inner._tbl.findall(f"{_NS}tr")
    # Проходим по строкам данных и строке ИТОГО: значений «… руб.» быть не должно.
    for tr in rows[1:]:
        for tc in tr.findall(f"{_NS}tc"):
            text = "".join(t.text or "" for t in tc.findall(f".//{_NS}t"))
            assert "руб" not in text.lower(), (
                f"В ячейке встретилось «руб»: {text!r}"
            )


def test_kp_itogo_filled_for_single_item():
    doc = _build_kp_with_fake([(50.0, 1, "ОдинТовар")])
    inner = _kp_inner_table(doc)
    rows = inner._tbl.findall(f"{_NS}tr")
    itogo_tc = rows[-1].findall(f"{_NS}tc")[-1]
    # 9А.2.5: разделитель тысяч теперь NBSP, нормализуем для сравнения.
    text = "".join(t.text or "" for t in itogo_tc.findall(f".//{_NS}t")).replace(
        " ", " "
    )
    assert text.strip(), "ИТОГО не должно быть пустым даже при одной позиции"
    # 50*90=4500 → +15% = 5175
    assert "5 175" in text


def test_kp_itogo_is_bold():
    doc = _build_kp_with_fake([(100.0, 1, "жирный итог")])
    inner = _kp_inner_table(doc)
    rows = inner._tbl.findall(f"{_NS}tr")
    itogo_tc = rows[-1].findall(f"{_NS}tc")[-1]
    # Найдём первый run и проверим rPr/b.
    run = itogo_tc.find(f".//{_NS}r")
    assert run is not None, "В ИТОГО нет run-а"
    rpr = run.find(f"{_NS}rPr")
    assert rpr is not None, "У run-а ИТОГО нет rPr — нельзя гарантировать bold"
    b_el = rpr.find(f"{_NS}b")
    assert b_el is not None, "В ИТОГО значение должно быть жирным (<w:b/>)"


# =============================================================================
# Этап 8.6 — KP: ИТОГО заполняется в правой нижней ячейке, ширины колонок,
# программно построенная таблица с 5 колонками
# =============================================================================

def test_kp_itogo_value_filled_in_last_cell():
    """Этап 8.6: значение ИТОГО — в последней ячейке последней строки.

    Раньше (после 8.4) шаблон содержал «лишние» gridCol-ы справа, и
    значение ИТОГО уходило в ячейку шириной 3545 twips за пределами
    видимой колонки «Сумма» — пользователь видел пустое поле.
    """
    doc = _build_kp_with_fake([(100.0, 2, "x")])
    inner = _kp_inner_table(doc)
    rows = inner._tbl.findall(f"{_NS}tr")
    itogo_tcs = rows[-1].findall(f"{_NS}tc")
    # В новой реализации ровно 2 ячейки: «ИТОГО» (gridSpan=4) + значение.
    assert len(itogo_tcs) == 2, (
        f"Ожидалось 2 ячейки в строке ИТОГО, получено {len(itogo_tcs)}"
    )
    value_tc = itogo_tcs[-1]
    text = "".join(t.text or "" for t in value_tc.findall(f".//{_NS}t")).replace(
        " ", " "
    )
    assert text.strip(), "Значение ИТОГО пустое"
    # 100*90=9000 → +15% = 10350 → ×2 = 20700.
    assert "20 700" in text, f"Ожидалось «20 700» в ИТОГО, получили {text!r}"


def test_kp_itogo_label_uses_grid_span_4():
    """Левая часть ИТОГО объединяет 4 колонки через gridSpan=4."""
    doc = _build_kp_with_fake([(50.0, 1, "single")])
    inner = _kp_inner_table(doc)
    rows = inner._tbl.findall(f"{_NS}tr")
    label_tc = rows[-1].findall(f"{_NS}tc")[0]
    tcPr = label_tc.find(f"{_NS}tcPr")
    grid_span = tcPr.find(f"{_NS}gridSpan") if tcPr is not None else None
    assert grid_span is not None and grid_span.get(f"{_NS}val") == "4", (
        "Ячейка «ИТОГО» должна объединять 4 колонки через gridSpan=4"
    )


def test_kp_table_columns_widths_for_price_and_sum():
    """Колонки «Цена» и «Сумма» — не уже 2.4 см (1361 twip).

    Иначе число вида «40 943» переносится на новую строку (живой баг 8.4).
    """
    doc = _build_kp_with_fake([(100.0, 1, "x")])
    inner = _kp_inner_table(doc)
    grid = inner._tbl.find(f"{_NS}tblGrid")
    cols = grid.findall(f"{_NS}gridCol")
    assert len(cols) == 5, f"Ожидалось 5 колонок в gridGrid, получено {len(cols)}"
    # cols: 0=№, 1=Наименование, 2=Кол-во, 3=Цена, 4=Сумма.
    price_w = int(cols[3].get(f"{_NS}w"))
    sum_w   = int(cols[4].get(f"{_NS}w"))
    # 2.4 см = 1361 twips (1 см ≈ 567 twips).
    assert price_w >= 1361, f"Колонка «Цена» {price_w} twips < 2.4см"
    assert sum_w   >= 1361, f"Колонка «Сумма» {sum_w} twips < 2.4см"


def test_kp_table_has_visible_borders():
    """Таблица КП имеет тонкие границы single со всех сторон."""
    doc = _build_kp_with_fake([(100.0, 1, "x")])
    inner = _kp_inner_table(doc)
    tbl_pr = inner._tbl.find(f"{_NS}tblPr")
    borders = tbl_pr.find(f"{_NS}tblBorders")
    assert borders is not None, "В tblPr должны быть tblBorders"
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.find(f"{_NS}{side}")
        assert b is not None, f"Граница {side} отсутствует"
        assert b.get(f"{_NS}val") == "single", (
            f"Граница {side}: ожидался single, {b.get(f'{_NS}val')!r}"
        )


# =============================================================================
# Блок D — GPU naming
# =============================================================================

def test_gpu_geforce_210_is_no_longer_rtx_1():
    raw = (
        "Видеокарта Biostar PCI-E G210-1GB D3 LP NVIDIA GeForce 210 "
        "1Gb 64bit DDR3 589/1333 DVIx1 HDMIx1 CRTx1 Ret low profile [VN2103NHG6]"
    )
    assert _short_gpu_model(raw) == "GeForce 210"


def test_gpu_rtx_not_matched_inside_word():
    # CRTx1 не должен давать ложный «RTX».
    raw = "CRTx1 Ret low profile"
    # В такой строке нет валидных маркеров → возвращаем то что осталось.
    out = _short_gpu_model(raw)
    assert out != "RTX x1" and not (out or "").startswith("RTX ")


def test_gpu_fallback_uses_brand_and_vram_on_suspicious_model():
    # Симулируем ситуацию, когда парсер всё же вернул подозрительное.
    out = _block_gpu(
        # Модель, из которой парсер извлечёт «RTX 1».
        "test something RTX 1 etc",
        {"vram_gb": 2, "vram_type": "DDR3"},
        "NVIDIA Corporation",
    )
    assert out is not None
    # В fallback нет «RTX 1», есть бренд и объём.
    assert "RTX 1" not in out
    assert "2GB" in out


def test_gpu_suspicious_detector():
    assert _looks_suspicious_gpu_model("RTX 1")
    assert _looks_suspicious_gpu_model("GTX 2")
    assert _looks_suspicious_gpu_model("Radeon 9")
    assert not _looks_suspicious_gpu_model("RTX 4060")
    assert not _looks_suspicious_gpu_model("Radeon RX 7600")
    assert not _looks_suspicious_gpu_model(None)
