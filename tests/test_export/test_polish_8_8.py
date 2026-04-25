# Юнит-тесты правок этапа 8.8: фикс паразитных gridCol во внешней
# обёрточной таблице KP, vAlign=center у всех ячеек, расширение
# колонки «Кол-во» с 1.2 → 1.4 см, корректный рендер quantity.
#
# Исходный баг: kp_template.docx содержит внешнюю таблицу 1×2 с
# gridCol [10490, 284], tblW=10774 и tblInd=-743. Сумма gridCol уезжает
# за полезную ширину A4 (9072 twips), а отрицательный отступ слева
# смещает таблицу за поля. Word при этом дезорганизует раскладку
# внутренней таблицы: «Кол-во» обрезается до «ол-во», quantity рендерится
# как «0», числа Цена/Сумма переносятся на вторую строку.

from __future__ import annotations

from datetime import date
from decimal import Decimal
from io import BytesIO
from unittest.mock import patch

import docx as _docx

from app.services.export import kp_builder


_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _fake_items(items_spec):
    out = []
    for i, (uu, qty, name) in enumerate(items_spec, start=1):
        out.append({
            "id": i, "query_id": 100 + i, "variant_manufacturer": "Intel",
            "quantity": qty, "position": i,
            "auto_name": name, "custom_name": None, "display_name": name,
            "unit_usd": uu, "unit_rub": 0.0,
            "total_usd": round(uu * qty, 2), "total_rub": 0.0,
            "created_at": None, "updated_at": None,
        })
    return out


def _mock_rate(value: str = "90"):
    return patch(
        "app.services.export.kp_builder.exchange_rate.get_usd_rate",
        return_value=(Decimal(value), date(2026, 4, 25), "cache"),
    )


def _build_kp(items_spec, markup=15):
    with patch(
        "app.services.export.kp_builder.spec_service.list_spec_items",
        return_value=_fake_items(items_spec),
    ), _mock_rate("90"):
        data = kp_builder.build_kp_docx(
            project_id=1, markup_percent=markup, db=None,
        )
    return _docx.Document(BytesIO(data))


def _all_tbls(doc):
    """Все <w:tbl> в document.xml, включая вложенные."""
    body = doc.element.body
    return body.findall(f".//{_NS}tbl")


def _inner_tbl(doc):
    """Внутренняя КП-таблица (та, что без вложенных таблиц)."""
    return doc.tables[0].rows[0].cells[0].tables[0]._tbl


def _outer_tbl(doc):
    return doc.tables[0]._tbl


# =============================================================================
# Inner-таблица: ровно 5 gridCol, tblW = сумма gridCol
# =============================================================================

def test_kp_no_phantom_gridcols():
    """Во внутренней таблице ровно 5 <w:gridCol> — по числу колонок
    (№, Наименование, Кол-во, Цена, Сумма). Никаких лишних остатков."""
    doc = _build_kp([(100.0, 2, "Тест")])
    inner = _inner_tbl(doc)
    grid = inner.find(f"{_NS}tblGrid")
    cols = grid.findall(f"{_NS}gridCol")
    assert len(cols) == 5, (
        f"Ожидалось 5 gridCol во внутренней таблице, получено {len(cols)}: "
        f"{[c.get(f'{_NS}w') for c in cols]}"
    )
    widths = [int(c.get(f"{_NS}w")) for c in cols]
    # Никаких подозрительных «обёрточных» ширин.
    for w in widths:
        assert 100 < w < 6000, (
            f"gridCol w={w} выглядит как остаток от внешней таблицы "
            f"(ожидалось значение в диапазоне 100..6000 twips)"
        )


def test_kp_tblw_matches_grid_sum():
    """tblW во внутренней таблице ровно равна сумме gridCol — иначе
    Word начинает auto-распределять ячейки и шапка «Кол-во»
    превращается в «ол-во»."""
    doc = _build_kp([(100.0, 1, "x")])
    inner = _inner_tbl(doc)
    grid = inner.find(f"{_NS}tblGrid")
    grid_sum = sum(int(gc.get(f"{_NS}w")) for gc in grid.findall(f"{_NS}gridCol"))
    tbl_w = inner.find(f"{_NS}tblPr").find(f"{_NS}tblW")
    assert int(tbl_w.get(f"{_NS}w")) == grid_sum, (
        f"tblW {tbl_w.get(f'{_NS}w')!r} != сумме gridCol {grid_sum}"
    )


# =============================================================================
# Outer-таблица: тоже нормализована, без отрицательного tblInd
# =============================================================================

def test_kp_no_negative_indent():
    """Ни в одной таблице document.xml не должно быть <w:tblInd> с
    отрицательным значением. Отрицательный indent сдвигает обёрточную
    таблицу за левое поле и ломает раскладку Word."""
    doc = _build_kp([(100.0, 1, "x")])
    for tbl in _all_tbls(doc):
        tblPr = tbl.find(f"{_NS}tblPr")
        if tblPr is None:
            continue
        tblInd = tblPr.find(f"{_NS}tblInd")
        if tblInd is None:
            continue
        val = int(tblInd.get(f"{_NS}w"))
        assert val >= 0, f"Найден отрицательный tblInd: {val}"


def test_kp_outer_table_normalized_to_inner_width():
    """Внешняя обёрточная таблица после нормализации шире внутренней быть
    не должна — иначе общая ширина уедет за полезные 9072 twips A4."""
    doc = _build_kp([(100.0, 1, "x")])
    outer = _outer_tbl(doc)
    grid = outer.find(f"{_NS}tblGrid")
    outer_sum = sum(int(gc.get(f"{_NS}w")) for gc in grid.findall(f"{_NS}gridCol"))
    assert outer_sum <= 9100, (
        f"Внешняя таблица шире 9072 twips: {outer_sum} (паразитные gridCol "
        f"не вычистили)"
    )
    # tblW должна совпадать с суммой gridCol.
    tbl_w = outer.find(f"{_NS}tblPr").find(f"{_NS}tblW")
    assert int(tbl_w.get(f"{_NS}w")) == outer_sum


# =============================================================================
# Структура строк: 5 ячеек в data, 2 в ИТОГО (gridSpan=4)
# =============================================================================

def test_kp_rows_have_expected_cell_counts():
    """В шапке и data-строках ровно 5 <w:tc>; в ИТОГО — 2 (label с
    gridSpan=4 + значение)."""
    doc = _build_kp([(100.0, 1, "Один"), (200.0, 3, "Три")])
    inner = _inner_tbl(doc)
    rows = inner.findall(f"{_NS}tr")
    # rows[0] — header, rows[1..-2] — data, rows[-1] — итог.
    assert len(rows) == 4  # header + 2 data + total
    assert len(rows[0].findall(f"{_NS}tc")) == 5
    assert len(rows[1].findall(f"{_NS}tc")) == 5
    assert len(rows[2].findall(f"{_NS}tc")) == 5
    total_tcs = rows[-1].findall(f"{_NS}tc")
    assert len(total_tcs) == 2
    gs = total_tcs[0].find(f"{_NS}tcPr").find(f"{_NS}gridSpan")
    assert gs is not None and gs.get(f"{_NS}val") == "4"


# =============================================================================
# vAlign center у всех ячеек таблицы
# =============================================================================

def test_kp_vertical_alignment_center():
    """В каждой ячейке внутренней таблицы (header + data + total)
    в tcPr задан <w:vAlign w:val="center"/>. Без этого числа в data-
    строках смотрятся прибитыми к верху и не совпадают по высоте
    с шапкой/ИТОГО.
    """
    doc = _build_kp([(100.0, 1, "x"), (50.0, 2, "y")])
    inner = _inner_tbl(doc)
    for tr_idx, tr in enumerate(inner.findall(f"{_NS}tr")):
        for c_idx, tc in enumerate(tr.findall(f"{_NS}tc")):
            tcPr = tc.find(f"{_NS}tcPr")
            assert tcPr is not None, f"row {tr_idx} cell {c_idx}: нет tcPr"
            vAlign = tcPr.find(f"{_NS}vAlign")
            assert vAlign is not None, (
                f"row {tr_idx} cell {c_idx}: нет <w:vAlign/>"
            )
            assert vAlign.get(f"{_NS}val") == "center", (
                f"row {tr_idx} cell {c_idx}: vAlign={vAlign.get(f'{_NS}val')!r}, "
                f"ожидался 'center'"
            )


# =============================================================================
# Quantity рендерится как реальное число
# =============================================================================

def test_kp_count_value_renders():
    """В data-строке колонка «Кол-во» (tc[2]) содержит реальное число
    quantity, а не «0». Если этот тест падает с «0» — баг не в рендере,
    а в том, что qty не сохранён в БД (или передан как 0/None)."""
    doc = _build_kp([(100.0, 10, "ДесятьШтук")])
    inner = _inner_tbl(doc)
    rows = inner.findall(f"{_NS}tr")
    # rows[0] header, rows[1] — единственная data-строка.
    qty_tc = rows[1].findall(f"{_NS}tc")[2]
    text = "".join(t.text or "" for t in qty_tc.findall(f".//{_NS}t"))
    assert text.strip() == "10", (
        f"Ожидалось «10» в колонке Кол-во, получили {text!r}"
    )


def test_kp_count_column_is_at_least_1_4_cm():
    """Колонка Кол-во ≥ 1.4 см (794 twips), чтобы заголовок «Кол-во»
    в 10pt не обрезался Word'ом до «ол-во»."""
    doc = _build_kp([(100.0, 1, "x")])
    inner = _inner_tbl(doc)
    grid = inner.find(f"{_NS}tblGrid")
    cols = grid.findall(f"{_NS}gridCol")
    qty_w = int(cols[2].get(f"{_NS}w"))
    assert qty_w >= 794, (
        f"Кол-во шириной {qty_w} twips < 794 (1.4см) — заголовок обрежется"
    )


# =============================================================================
# noWrap не пропал при переключении на center и расширении Кол-во
# =============================================================================

def test_kp_price_sum_cells_keep_no_wrap():
    """Колонки Цена (tc[3]) и Сумма (tc[4]) во всех data-строках всё
    ещё имеют <w:noWrap/> — этап 8.7 не должен сломаться."""
    doc = _build_kp([(100.0, 1, "x"), (200.0, 3, "y")])
    inner = _inner_tbl(doc)
    rows = inner.findall(f"{_NS}tr")
    for tr_idx, tr in enumerate(rows[1:-1], start=1):
        tcs = tr.findall(f"{_NS}tc")
        for col_idx in (3, 4):
            tcPr = tcs[col_idx].find(f"{_NS}tcPr")
            no_wrap = tcPr.find(f"{_NS}noWrap")
            assert no_wrap is not None, (
                f"row {tr_idx} cell {col_idx}: noWrap пропал после правок 8.8"
            )
